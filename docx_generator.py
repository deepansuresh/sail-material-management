import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def generate_purchase_proposal_docx(data: dict, output_path: str):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(30, 41, 59)

    # Title / Header
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("PURCHASE PROPOSAL NOTE")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("STEEL AUTHORITY OF INDIA LIMITED - SALEM STEEL PLANT")
    r_sub.bold = True
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Main Indent Particulars Table (Matching Screenshot 2 & 3)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header Row: Description of the item
    row0 = table.rows[0]
    cell_top = row0.cells[0]
    row0.cells[0].merge(row0.cells[1])
    p = cell_top.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run_bold = p.add_run(f"Description of the item: {data.get('item_description', '')}")
    run_bold.bold = True
    set_cell_background(cell_top, "F1F5F9")

    def add_section_header(tbl, title):
        r = tbl.add_row()
        c = r.cells[0]
        r.cells[0].merge(r.cells[1])
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title)
        run.bold = True
        set_cell_background(c, "E2E8F0")

    def add_key_value(tbl, k, v):
        r = tbl.add_row()
        c1, c2 = r.cells[0], r.cells[1]
        c1.paragraphs[0].add_run(k).bold = True
        c2.paragraphs[0].add_run(str(v))
        c1.width = Inches(2.8)
        c2.width = Inches(4.3)
        set_cell_background(c1, "F8FAFC")

    # Indent Particulars
    add_section_header(table, "Indent Particulars")
    ind = data.get("indent_particulars", {})
    add_key_value(table, "Purchase requisition no.", ind.get("purchase_requisition_no", ""))
    add_key_value(table, "Indent reference no.", ind.get("indent_reference_no", ""))
    add_key_value(table, "Indent date", ind.get("indent_date", ""))
    add_key_value(table, "Proposal date", ind.get("proposal_date", ""))
    add_key_value(table, "Indent raised by", ind.get("indent_raised_by", ""))
    add_key_value(table, "Estimate", ind.get("estimate", ""))
    add_key_value(table, "Basis of estimate", ind.get("basis_of_estimate", ""))
    add_key_value(table, "First time procurement", ind.get("first_time_procurement", ""))
    add_key_value(table, "Number of budgetary offers received", ind.get("budgetary_offers_count", ""))

    # Previous purchase details
    add_section_header(table, "Previous purchase details:")
    prev = data.get("previous_purchase_details", {})
    
    # Nested table or rows for previous purchases
    r_prev = table.add_row()
    c_prev = r_prev.cells[0]
    r_prev.cells[0].merge(r_prev.cells[1])
    p_p = c_prev.paragraphs[0]
    
    # Add embedded sub-table
    sub_tbl = c_prev.add_table(rows=1, cols=4)
    set_table_borders(sub_tbl, "CBD5E1")
    headers = ["Item sl. nos.", "AT ref. no.", "Previous purchase qty", "Unit rate incl. GST"]
    for i, h in enumerate(headers):
        cell = sub_tbl.rows[0].cells[i]
        cell.paragraphs[0].add_run(h).bold = True
        set_cell_background(cell, "F1F5F9")
    
    for itm in prev.get("items", []):
        r_sub = sub_tbl.add_row()
        r_sub.cells[0].paragraphs[0].add_run(str(itm.get("item_sl_no", "")))
        r_sub.cells[1].paragraphs[0].add_run(str(itm.get("at_ref_no", "")))
        r_sub.cells[2].paragraphs[0].add_run(str(itm.get("prev_qty", "")))
        r_sub.cells[3].paragraphs[0].add_run(str(itm.get("unit_rate_incl_gst", "")))

    add_key_value(table, "Previous purchase mode of tender", prev.get("prev_mode_of_tender", ""))

    # Indent Approval
    add_section_header(table, "Indent Approval")
    ia = data.get("indent_approval", {})
    add_key_value(table, "Approving Authority", ia.get("approving_authority", ""))
    add_key_value(table, "Indent approved date", ia.get("indent_approved_date", ""))
    add_key_value(table, "Mode of Tender", ia.get("mode_of_tender", ""))

    # Sanction Particulars
    add_section_header(table, "Sanction Particulars:")
    sp = data.get("sanction_particulars", {})
    add_key_value(table, "Name of the supplier", sp.get("supplier_name", ""))
    add_key_value(table, "Order Value Incl. GST", sp.get("order_value_incl_gst", ""))
    add_key_value(table, "Deviation wrt estimate", sp.get("deviation_wrt_estimate", ""))

    # Negotiation Details
    add_section_header(table, "Negotiation Details")
    neg = data.get("negotiation_details", {})
    
    r_neg = table.add_row()
    c_neg = r_neg.cells[0]
    r_neg.cells[0].merge(r_neg.cells[1])
    
    neg_tbl = c_neg.add_table(rows=1, cols=3)
    set_table_borders(neg_tbl, "CBD5E1")
    n_headers = neg.get("headers", ["Parameter", "Tender Price", "After Negotiation"])
    for i, h in enumerate(n_headers):
        cell = neg_tbl.rows[0].cells[i]
        cell.paragraphs[0].add_run(h).bold = True
        set_cell_background(cell, "F1F5F9")
        
    for row in neg.get("rows", []):
        r_n = neg_tbl.add_row()
        r_n.cells[0].paragraphs[0].add_run(str(row[0])).bold = True
        r_n.cells[1].paragraphs[0].add_run(str(row[1]))
        r_n.cells[2].paragraphs[0].add_run(str(row[2]) if len(row) > 2 else "")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Narrative Clauses 1 to 9 (Matching Screenshot 3 & 4)
    for i, clause in enumerate(data.get("narrative_clauses", []), 1):
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_after = Pt(6)
        p_c.paragraph_format.line_spacing = 1.15
        
        # Check if clause already starts with number
        prefix = f"{i}. " if not clause.strip().startswith(str(i)) else ""
        run = p_c.add_run(prefix + clause)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Proposed Order Terms Table (Screenshot 4)
    pot = data.get("proposed_order_terms", {})
    tbl_terms = doc.add_table(rows=0, cols=2)
    tbl_terms.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_terms)

    terms_kv = [
        ("Supplier Name", pot.get("supplier_name", "")),
        ("Item Description", pot.get("item_description", "")),
        ("Total Order Value without GST", pot.get("total_order_value_without_gst", "")),
        ("Total Order Value with GST", pot.get("total_order_value_with_gst", "")),
        ("Estimate", pot.get("estimate", "")),
        ("% Dev wrt estimate", pot.get("percent_dev_wrt_estimate", "")),
    ]
    for k, v in terms_kv:
        add_key_value(tbl_terms, k, v)

    # Commercial Terms sub-section
    add_section_header(tbl_terms, "Commercial Terms:")
    ct = pot.get("commercial_terms", {})
    add_key_value(tbl_terms, "Terms of Delivery", ct.get("terms_of_delivery", ""))
    add_key_value(tbl_terms, "Delivery Schedule", ct.get("delivery_schedule", ""))
    add_key_value(tbl_terms, "Payment Terms", ct.get("payment_terms", ""))
    add_key_value(tbl_terms, "Offer Validity", ct.get("offer_validity", ""))

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Approval Sought For
    p_as = doc.add_paragraph()
    p_as.add_run("Approval Sought For").bold = True
    p_as.paragraph_format.space_after = Pt(3)
    p_as_text = doc.add_paragraph(data.get("approval_sought_for", ""))
    p_as_text.paragraph_format.space_after = Pt(10)

    # Approving Authority / DOP / Manual Ref
    p_dop = doc.add_paragraph()
    p_dop.add_run("Approving Authority / DOP / Manual Ref").bold = True
    p_dop.paragraph_format.space_after = Pt(3)
    p_dop_text = doc.add_paragraph(data.get("approving_authority_dop", ""))
    p_dop_text.paragraph_format.space_after = Pt(10)

    # Suggested Approval Path
    p_ap = doc.add_paragraph()
    p_ap.add_run("Suggested Approval Path").bold = True
    p_ap.paragraph_format.space_after = Pt(3)
    p_ap_text = doc.add_paragraph(data.get("suggested_approval_path", ""))
    p_ap_text.paragraph_format.space_after = Pt(14)

    doc.save(output_path)
